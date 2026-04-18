"""
PDF and DOCX text parser.
"""
import io

import fitz  # PyMuPDF
import pdfplumber
from docx import Document


def parse_file(file_storage) -> str:
    """
    Accept a Flask FileStorage object, detect type, return raw text.
    """
    filename = file_storage.filename.lower()
    raw_bytes = file_storage.read()

    if filename.endswith(".pdf"):
        return _parse_pdf(raw_bytes)
    elif filename.endswith((".docx", ".doc")):
        return _parse_docx(raw_bytes)
    else:
        raise ValueError("Unsupported file type. Upload a PDF or DOCX.")


def _parse_pdf(raw_bytes: bytes) -> str:
    """Try PyMuPDF first; fall back to pdfplumber for scanned/tricky PDFs."""
    text = ""
    try:
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception:
        text = ""

    # Fallback: pdfplumber
    if len(text.strip()) < 100:
        try:
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception:
            pass

    return text.strip()


def _parse_docx(raw_bytes: bytes) -> str:
    doc = Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
